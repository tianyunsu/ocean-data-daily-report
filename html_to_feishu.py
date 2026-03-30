#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
html_to_feishu.py
1. 解析本地HTML日报 → 生成结构化Word文档(.docx)
2. 上传HTML原文件 + docx 到飞书云盘
3. 通过飞书import_task API将docx导入为原生飞书文档
"""

import os, time, json, requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 配置 ────────────────────────────────────────────────
APP_ID     = "cli_a93d483f6ff81bca"
APP_SECRET = "CU3EPesfCzNayK4bqsnh6droaJsf4HV8"
TENANT_DOMAIN = "wcn5jx0ifkx3.feishu.cn"

HTML_FILE  = r"C:\Users\Administrator\WorkBuddy\Claw\daily_reports\海洋AI简报_2026-03-16.html"
DOCX_FILE  = r"C:\Users\Administrator\WorkBuddy\Claw\daily_reports\海洋AI简报_2026-03-16.docx"

# ─── 获取 token ──────────────────────────────────────────
def get_token():
    r = requests.post(
        "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
        json={"app_id": APP_ID, "app_secret": APP_SECRET}, timeout=15
    )
    return r.json()["tenant_access_token"]

# ─── Step 1: 解析HTML → 生成Word ─────────────────────────
def parse_html(html_path):
    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    report = {}
    # 标题
    h1 = soup.find("h1")
    report["title"] = h1.get_text(strip=True) if h1 else "海洋AI技术日报"

    # 统计行
    stats = []
    for item in soup.select(".stat-item"):
        num   = item.select_one(".stat-num")
        label = item.select_one(".stat-label")
        if num and label:
            stats.append(f"{num.get_text(strip=True)} {label.get_text(strip=True)}")
    report["stats"] = stats

    # 各板块
    sections = []
    for sec in soup.select(".section"):
        sec_title = sec.select_one(".section-title")
        sec_en    = sec.select_one(".section-en")
        sec_count = sec.select_one(".section-count")
        articles  = []
        for idx, art in enumerate(sec.select(".article-item"), 1):
            title_el  = art.select_one(".article-title")
            badge_el  = art.select_one(".highlight-badge")
            abstract  = art.select_one(".article-abstract")
            source    = art.select_one(".tag-source")
            date_tag  = art.select_one(".tag-date")
            link_tag  = art.select_one(".tag-link")

            badge_text = badge_el.get_text(strip=True) if badge_el else ""
            # 标题去掉badge文字
            title_text = title_el.get_text(strip=True) if title_el else ""
            if badge_text and title_text.startswith(badge_text):
                title_text = title_text[len(badge_text):].strip()

            articles.append({
                "index":    idx,
                "badge":    badge_text,
                "title":    title_text,
                "abstract": abstract.get_text(" ", strip=True) if abstract else "",
                "source":   source.get_text(strip=True) if source else "",
                "date":     date_tag.get_text(strip=True) if date_tag else "",
                "link":     link_tag.get("href", "") if link_tag else "",
            })
        sections.append({
            "title":   sec_title.get_text(strip=True) if sec_title else "",
            "en":      sec_en.get_text(strip=True) if sec_en else "",
            "count":   sec_count.get_text(strip=True) if sec_count else "",
            "articles": articles,
        })
    report["sections"] = sections

    # footer
    footer = soup.select_one(".footer")
    report["footer"] = footer.get_text(" ", strip=True) if footer else ""
    return report


def set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def add_heading(doc, text, level=1, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = color or RGBColor(0x00, 0x5b, 0x9a)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return para


def build_docx(report, out_path):
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── 标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run(report["title"])
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # ── 副标题 ──
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(4)
    sub_run = sub_para.add_run("汇聚全球海洋人工智能领域最新研究动态 | 每日智能推送")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x5a, 0x6a, 0x7a)

    # ── 统计行 ──
    if report.get("stats"):
        stats_para = doc.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stats_para.paragraph_format.space_after = Pt(8)
        stats_run = stats_para.add_run("  |  ".join(report["stats"]))
        stats_run.font.size = Pt(10)
        stats_run.bold = True
        stats_run.font.color.rgb = RGBColor(0x00, 0xa8, 0xb5)

    # 分隔线
    doc.add_paragraph("─" * 60)

    section_colors = [
        RGBColor(0x00, 0x3d, 0x80),
        RGBColor(0x00, 0x4d, 0x6e),
        RGBColor(0x00, 0x50, 0x60),
        RGBColor(0x00, 0x3d, 0x5c),
        RGBColor(0x00, 0x2d, 0x4d),
    ]

    for s_idx, sec in enumerate(report["sections"]):
        color = section_colors[s_idx % len(section_colors)]

        # ── 板块标题 ──
        sec_para = doc.add_paragraph()
        sec_para.paragraph_format.space_before = Pt(14)
        sec_para.paragraph_format.space_after  = Pt(2)
        icon_map = {"海洋人工智能": "🤖", "海洋数字孪生": "🌐",
                    "海洋可视化": "📊", "海洋数据质量": "✅", "海洋数据处理": "⚙️"}
        icon = icon_map.get(sec["title"], "●")
        run = sec_para.add_run(f"{icon}  {sec['title']}  ({sec['count']})")
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = color

        # 英文子标题
        en_para = doc.add_paragraph()
        en_para.paragraph_format.space_before = Pt(0)
        en_para.paragraph_format.space_after  = Pt(6)
        en_run = en_para.add_run(sec["en"])
        en_run.font.size = Pt(9)
        en_run.font.color.rgb = RGBColor(0x5a, 0x6a, 0x7a)
        en_run.italic = True

        # ── 各条资讯 ──
        for art in sec["articles"]:
            # 标题行
            art_para = doc.add_paragraph()
            art_para.paragraph_format.space_before = Pt(8)
            art_para.paragraph_format.space_after  = Pt(2)

            # 序号
            num_run = art_para.add_run(f"{art['index']}. ")
            num_run.bold = True
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = color

            # badge
            if art["badge"]:
                badge_run = art_para.add_run(f"[{art['badge']}] ")
                badge_run.bold = True
                badge_run.font.size = Pt(11)
                badge_run.font.color.rgb = RGBColor(0xf0, 0xa5, 0x00)

            # 正标题
            title_run = art_para.add_run(art["title"])
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

            # 摘要
            if art["abstract"]:
                abs_para = doc.add_paragraph()
                abs_para.paragraph_format.space_before = Pt(2)
                abs_para.paragraph_format.space_after  = Pt(2)
                abs_para.paragraph_format.left_indent  = Cm(0.8)
                abs_run = abs_para.add_run(art["abstract"])
                abs_run.font.size = Pt(10)
                abs_run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x3a)

            # 元信息行
            meta_parts = []
            if art["source"]:
                meta_parts.append(art["source"])
            if art["date"]:
                meta_parts.append(art["date"])
            if art["link"]:
                meta_parts.append(f"链接: {art['link']}")
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_para.paragraph_format.space_before = Pt(1)
                meta_para.paragraph_format.space_after  = Pt(4)
                meta_para.paragraph_format.left_indent  = Cm(0.8)
                meta_run = meta_para.add_run("  |  ".join(meta_parts))
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(0x5a, 0x6a, 0x7a)

        # 板块分隔
        doc.add_paragraph()

    # ── 页脚 ──
    doc.add_paragraph("─" * 60)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(report.get("footer", ""))
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x5a, 0x6a, 0x7a)

    doc.save(out_path)
    print(f"[OK] Word文档已生成: {out_path}")


# ─── Step 2: 上传文件到飞书云盘 ──────────────────────────
def upload_to_drive(token, file_path, file_name, parent_token=None, parent_type="explorer"):
    """上传文件到飞书云盘，返回 file_token"""
    file_size = os.path.getsize(file_path)

    # 1. 预上传
    pre_url = "https://open.feishu.cn/open-apis/drive/v1/files/upload_prepare"
    pre_body = {
        "file_name": file_name,
        "parent_type": parent_type,
        "size": file_size,
    }
    if parent_token:
        pre_body["parent_node"] = parent_token

    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    pre_r = requests.post(pre_url, json=pre_body, headers=headers, timeout=15)
    pre_data = pre_r.json()
    print(f"  预上传响应: {json.dumps(pre_data, ensure_ascii=False)[:300]}")

    if pre_data.get("code") != 0:
        print(f"  [ERROR] 预上传失败: {pre_data}")
        return None

    upload_id  = pre_data["data"]["upload_id"]
    block_size = pre_data["data"]["block_size"]
    block_num  = pre_data["data"]["block_num"]

    # 2. 分块上传
    with open(file_path, "rb") as f:
        for block_idx in range(block_num):
            chunk = f.read(block_size)
            part_url = "https://open.feishu.cn/open-apis/drive/v1/files/upload_part"
            part_headers = {"Authorization": f"Bearer {token}"}
            part_r = requests.post(
                part_url,
                headers=part_headers,
                data={"upload_id": upload_id, "seq": block_idx},
                files={"file": (file_name, chunk)},
                timeout=30,
            )
            part_data = part_r.json()
            if part_data.get("code") != 0:
                print(f"  [ERROR] 分块{block_idx}上传失败: {part_data}")
                return None
            print(f"  分块 {block_idx+1}/{block_num} 上传成功")

    # 3. 完成上传
    finish_url = "https://open.feishu.cn/open-apis/drive/v1/files/upload_finish"
    finish_r = requests.post(
        finish_url,
        json={"upload_id": upload_id, "block_num": block_num},
        headers=headers,
        timeout=15,
    )
    finish_data = finish_r.json()
    print(f"  完成上传响应: {json.dumps(finish_data, ensure_ascii=False)[:300]}")

    if finish_data.get("code") == 0:
        file_token = finish_data["data"]["file_token"]
        print(f"  [OK] 上传成功，file_token: {file_token}")
        return file_token
    else:
        print(f"  [ERROR] 完成上传失败: {finish_data}")
        return None


# ─── Step 3: 创建导入任务 ────────────────────────────────
def import_docx_to_doc(token, file_token, file_name, folder_token=None):
    """将docx文件导入为飞书原生文档，返回 (ticket, import_token)"""
    url = "https://open.feishu.cn/open-apis/drive/v1/import_tasks"
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    body = {
        "file_extension": "docx",
        "file_token": file_token,
        "type": "docx",
        "file_name": file_name.replace(".docx", ""),
    }
    if folder_token:
        body["point"] = {"mount_type": 1, "mount_key": folder_token}

    r = requests.post(url, json=body, headers=headers, timeout=15)
    data = r.json()
    print(f"  导入任务响应: {json.dumps(data, ensure_ascii=False)[:400]}")

    if data.get("code") == 0:
        ticket = data["data"]["ticket"]
        print(f"  [OK] 导入任务已创建，ticket: {ticket}")
        return ticket
    else:
        print(f"  [ERROR] 创建导入任务失败: {data}")
        return None


def poll_import_result(token, ticket, max_wait=60):
    """轮询导入任务结果，返回文档token"""
    url = f"https://open.feishu.cn/open-apis/drive/v1/import_tasks/{ticket}"
    headers = {"Authorization": f"Bearer {token}"}
    for i in range(max_wait // 3):
        time.sleep(3)
        r = requests.get(url, headers=headers, timeout=15)
        data = r.json()
        if data.get("code") != 0:
            print(f"  轮询失败: {data}")
            break
        result = data["data"]["result"]
        status = result.get("job_status")
        print(f"  [{i+1}] 导入状态: {status}")
        if status == 0:  # 成功
            doc_token = result.get("token")
            print(f"  [OK] 导入完成！文档token: {doc_token}")
            return doc_token
        elif status in (1, 2):  # 进行中
            continue
        else:
            job_error = result.get("job_error_msg", "未知错误")
            print(f"  [ERROR] 导入失败: status={status}, error={job_error}")
            break
    return None


# ─── 主流程 ──────────────────────────────────────────────
def main():
    print("=" * 60)
    print("Step 1: 解析HTML → 生成Word文档")
    print("=" * 60)
    report = parse_html(HTML_FILE)
    print(f"  标题: {report['title']}")
    print(f"  板块数: {len(report['sections'])}")
    for s in report["sections"]:
        print(f"    - {s['title']}: {len(s['articles'])} 条")

    build_docx(report, DOCX_FILE)

    print()
    print("=" * 60)
    print("Step 2: 获取飞书访问凭证")
    print("=" * 60)
    token = get_token()
    print(f"  [OK] token: {token[:16]}...")

    print()
    print("=" * 60)
    print("Step 3: 上传HTML文件到飞书云盘")
    print("=" * 60)
    html_token = upload_to_drive(token, HTML_FILE, "海洋AI简报_2026-03-16.html")

    print()
    print("=" * 60)
    print("Step 4: 上传Word文件到飞书云盘")
    print("=" * 60)
    docx_token = upload_to_drive(token, DOCX_FILE, "海洋AI简报_2026-03-16.docx")

    print()
    print("=" * 60)
    print("Step 5: 将Word导入为飞书原生文档")
    print("=" * 60)
    doc_url = None
    if docx_token:
        ticket = import_docx_to_doc(token, docx_token, "海洋AI简报_2026-03-16.docx")
        if ticket:
            print("  等待导入完成（最多60秒）...")
            doc_token = poll_import_result(token, ticket, max_wait=60)
            if doc_token:
                doc_url = f"https://{TENANT_DOMAIN}/docx/{doc_token}"
                print(f"\n  [OK] 原生文档链接: {doc_url}")

    print()
    print("=" * 60)
    print("全部完成！")
    print("=" * 60)
    if html_token:
        print(f"HTML云盘文件 token : {html_token}")
        print(f"HTML 云盘链接       : https://{TENANT_DOMAIN}/file/{html_token}")
    if docx_token:
        print(f"Word云盘文件 token : {docx_token}")
        print(f"Word 云盘链接       : https://{TENANT_DOMAIN}/file/{docx_token}")
    if doc_url:
        print(f"原生飞书文档链接    : {doc_url}")

    # 写入url文件
    if doc_url:
        with open(r"C:\Users\Administrator\WorkBuddy\Claw\feishu_doc_url.txt", "w", encoding="utf-8") as f:
            f.write(doc_url + "\n")


if __name__ == "__main__":
    main()
